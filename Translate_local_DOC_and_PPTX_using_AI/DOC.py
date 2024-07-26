from docx import Document


def process_paragraph(paragraph, process_function):
    if paragraph.text.strip():
        original_text = paragraph.text
        corrected_text = process_function(original_text)
        #print(f"Original: {original_text}\nTranslated: {corrected_text}\n")  # Debug statement
        paragraph.text = corrected_text


def process_table(table, process_function):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, process_function)


def process_docx(file_path, process_function, progress_callback=None):
    doc = Document(file_path)
    total_paragraphs = len(doc.paragraphs)
    # Procesează textul din document
    for index, paragraph in enumerate(doc.paragraphs):
        process_paragraph(paragraph, process_function)
        if progress_callback:
            progress = (index + 1) / total_paragraphs * 100
            progress_callback(progress)

    # Procesează tabelele din document
    for table in doc.tables:
        process_table(table, process_function)

    return doc